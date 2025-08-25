import re
import PyPDF2
import openpyxl
from docx import Document # Importa a biblioteca para trabalhar com .docx
from docx.shared import Inches # Para unidades de medida em .docx
import datetime # Para manipulação de datas

# ==============================================================================
# FUNÇÃO 1: Extrai o texto bruto do PDF
# ==============================================================================
def extrair_texto_de_pdf(caminho_do_pdf):
    """
    Extrai o texto de todas as páginas de um arquivo PDF.
    """
    texto_completo = ""
    try:
        with open(caminho_do_pdf, 'rb') as arquivo:
            leitor_pdf = PyPDF2.PdfReader(arquivo)
            for pagina in leitor_pdf.pages:
                texto_completo += pagina.extract_text() + "\n"
    except FileNotFoundError:
        print(f"ERRO: Arquivo não encontrado em '{caminho_do_pdf}'. Verifique o caminho.")
        return None
    except Exception as e:
        print(f"Ocorreu um erro inesperado ao ler o PDF: {e}")
        return None
    return texto_completo

# ==============================================================================
# FUNÇÃO 2: Extrai os dados específicos do texto
# ==============================================================================
def extrair_dados_do_contrato(texto_do_contrato):
    """
    Extrai dados específicos do conteúdo textual de um contrato.
    """
    dados_extraidos = {}
    flags = re.DOTALL | re.IGNORECASE

    # 1. Extrair dados do CONTRATANTE
    contratante_match = re.search(
        r"CONTRATANTE\s*:\s*Sr\(a\)\s*(.*?),\s*brasileiro\(a\).*?RG:\s*([\d.\s-]+?)\s*e\s*CPF:\s*([\d.\s-]+?),",
        texto_do_contrato, flags)
    if contratante_match:
        dados_extraidos['Contratante'] = {'Nome': contratante_match.group(1).strip(), 'RG': contratante_match.group(2).strip(), 'CPF': contratante_match.group(3).strip(),}
    else:
        dados_extraidos['Contratante'] = "Não encontrado"

    # 2. Extrair dados do CONTRATADO
    contratado_match = re.search(
        r"CONTRATADO\s*(.*?),\s*inscrito\s*sob\s*o\s*CNPJ:\s*([\d./\s-]+?),",
        texto_do_contrato, flags)
    if contratado_match:
        dados_extraidos['Contratado'] = {'Nome Empresa': contratado_match.group(1).strip(), 'CNPJ': contratado_match.group(2).strip(),}
    else:
        dados_extraidos['Contratado'] = "Não encontrado"

    # 3. Extrair Produtos Contratados
    dados_extraidos['Produtos Contratados'] = []
    secao_produtos_match = re.search(r'PRODUTOS CONTRATADOS\s*(.*?)\s*CLÁUSULA 2', texto_do_contrato, flags)
    if secao_produtos_match:
        texto_tabela = secao_produtos_match.group(1)
        linhas = texto_tabela.strip().split('\n')
        for linha in linhas:
            linha_limpa = linha.strip()
            if not linha_limpa or not linha_limpa[0].isdigit():
                continue
            # Ajuste a regex se o formato da tabela for diferente
            # Ex: (Quantidade) (Produto) (Valor Unitário) (Valor Total Item)
            produto_match = re.match(r'^\s*(\d+)\s+(.*?)\s+([\d,.]+)\s+([\d,.]+)\s*$', linha_limpa)
            if produto_match:
                dados_extraidos['Produtos Contratados'].append({
                    'Quantidade': produto_match.group(1).strip(),
                    'Produto': produto_match.group(2).strip().strip(),
                    'Valor Unitário': produto_match.group(3).strip(),
                    'Valor Total Item': produto_match.group(4).strip()
                })

    # 4, 5, 6. Extrair outros dados
    valor_total_match = re.search(r"O valor total de R\$\s*([\d,.]+)", texto_do_contrato, flags)
    dados_extraidos['Valor Total do Pedido'] = valor_total_match.group(1).strip() if valor_total_match else "Não encontrado"
    
    data_pagamento_match = re.search(r"pagos no dia\s*(\d{2}/\d{2}/\d{4})", texto_do_contrato, flags)
    dados_extraidos['Data de Pagamento'] = data_pagamento_match.group(1).strip() if data_pagamento_match else "Não encontrada"
    
    data_evento_match = re.search(r"O evento acontecerá no dia:\s*([\d/]+)", texto_do_contrato, flags)
    dados_extraidos['Data do Evento'] = data_evento_match.group(1).strip() if data_evento_match else "Não encontrada"
    
    local_evento_match = re.search(r"Local do\s*evento\s*:\s*(.*?)\n", texto_do_contrato, flags)
    dados_extraidos['Local do Evento'] = local_evento_match.group(1).strip() if local_evento_match else "Não encontrado"
            
    return dados_extraidos

# ==============================================================================
# FUNÇÃO 3: Exporta os dados para um arquivo Excel
# ==============================================================================
def exportar_para_excel(dados, nome_do_arquivo="dados_contrato.xlsx"):
    """
    Exporta o dicionário de dados extraídos para uma planilha Excel.
    Esta função é focada em dados de contrato (singular).
    Para exportar uma LISTA de pedidos, a lógica precisará ser adaptada.
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Dados do Contrato"

    sheet['A1'] = "Campo"
    sheet['B1'] = "Informação Extraída"
    
    linha_atual = 2

    for chave, valor in dados.items():
        if chave == 'Produtos Contratados':
            continue
        
        if isinstance(valor, dict):
            for sub_chave, sub_valor in valor.items():
                sheet[f'A{linha_atual}'] = f"{chave} - {sub_chave}"
                sheet[f'B{linha_atual}'] = sub_valor
                linha_atual += 1
        else:
            sheet[f'A{linha_atual}'] = chave
            sheet[f'B{linha_atual}'] = valor
            linha_atual += 1

    # Adiciona a tabela de produtos contratados
    linha_atual += 2 
    
    headers_produtos = list(dados['Produtos Contratados'][0].keys()) if dados.get('Produtos Contratados') and dados['Produtos Contratados'] else []
    if headers_produtos:
        for col_idx, header in enumerate(headers_produtos, 1):
            sheet.cell(row=linha_atual, column=col_idx, value=header)
        
        linha_atual += 1

        for produto in dados['Produtos Contratados']:
            for col_idx, header in enumerate(headers_produtos, 1):
                sheet.cell(row=linha_atual, column=col_idx, value=produto.get(header, 'N/A'))
            linha_atual += 1
    
    try:
        workbook.save(filename=nome_do_arquivo)
        print(f"\n[SUCESSO] Dados exportados para o arquivo '{nome_do_arquivo}'")
    except Exception as e:
        print(f"\n[ERRO] Não foi possível salvar a planilha: {e}")

# ==============================================================================
# FUNÇÃO 4: Gera o Relatório de Entrega em DOCX (MOVIMENTO DO relatorio_entrega.py)
# ==============================================================================
def gerar_relatorio_entrega(dados, nome_arquivo='relatorio_entrega.docx'):
    """
    Gera um relatório de entrega em formato .docx com base nos dados do pedido.
    """
    document = Document()
    
    # Título
    document.add_heading('RELATÓRIO DE ENTREGA', 0)
    
    # Informações principais
    contratante = dados.get('Contratante', {})
    data_evento = dados.get('Data do Evento', 'Não informada')
    local_evento = dados.get('Local do Evento', 'Não informado')
    
    document.add_paragraph(f"Nome do Cliente: {contratante.get('Nome', 'Não encontrado')}")
    document.add_paragraph(f"Data do Evento: {data_evento}")
    document.add_paragraph(f"Local do Evento: {local_evento}")
    document.add_paragraph(f"Data de Emissão: {datetime.datetime.now().strftime('%d/%m/%Y')}")

    document.add_paragraph("\nProdutos Contratados:")
    
    # Tabela de produtos
    produtos = dados.get('Produtos Contratados', [])
    if produtos:
        tabela = document.add_table(rows=1, cols=4)
        tabela.style = 'Table Grid'
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Quantidade'
        hdr_cells[1].text = 'Produto'
        hdr_cells[2].text = 'Valor Unitário'
        hdr_cells[3].text = 'Valor Total'

        for item in produtos:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(item.get('Quantidade', 'N/A')) # Garante string
            row_cells[1].text = str(item.get('Produto', 'N/A'))
            row_cells[2].text = str(item.get('Valor Unitário', 'N/A'))
            row_cells[3].text = str(item.get('Valor Total Item', 'N/A'))
    else:
        document.add_paragraph("Nenhum produto encontrado.")

    document.add_paragraph(f"\nValor Total do Pedido: R$ {dados.get('Valor Total do Pedido', 'Não encontrado')}")

    # Espaço para assinaturas
    document.add_paragraph("\n\n\nAssinaturas:\n")
    document.add_paragraph("______________________________\nResponsável pela Entrega")
    document.add_paragraph("\n\n")
    document.add_paragraph("______________________________\nResponsável pela Retirada")

    # Salvar o arquivo
    try:
        document.save(nome_arquivo)
        print(f"[OK] Relatório de entrega salvo em: {nome_arquivo}")
    except Exception as e:
        print(f"[ERRO] Falha ao salvar relatório de entrega: {e}")

# ==============================================================================
# BLOCO DE EXECUÇÃO: O "Gerente de Operações" para testes local
# ==============================================================================
if __name__ == "__main__":
    caminho_do_pdf = "modelo_contrato.pdf" # Certifique-se de ter este arquivo no mesmo diretório

    # Exemplo de uso das funções
    print(f"Tentando extrair texto de: {caminho_do_pdf}")
    texto_extraido = extrair_texto_de_pdf(caminho_do_pdf)
    
    if texto_extraido:
        print("\nTexto extraído com sucesso. Tentando extrair dados...")
        dados_do_contrato = extrair_dados_do_contrato(texto_extraido)
        
        # Imprime os resultados no terminal
        print("\n--- DADOS EXTRAÍDOS DO CONTRATO ---")
        for chave, valor in dados_do_contrato.items():
            print(f"\n>> {chave}:")
            if isinstance(valor, dict):
                for sub_chave, sub_valor in valor.items():
                    print(f"    {sub_chave}: {sub_valor}")
            elif isinstance(valor, list) and valor:
                for item in valor:
                    print(f"    - {item}")
            else:
                print(f"    {valor}")
        print("\n----------------------------------------------------")

        # Chama a função para criar a planilha Excel
        nome_excel = f"dados_contrato_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        exportar_para_excel(dados_do_contrato, nome_do_arquivo=nome_excel)

        # Chama a nova função para criar o relatório de entrega em DOCX
        nome_docx = f"relatorio_entrega_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        gerar_relatorio_entrega(dados_do_contrato, nome_arquivo=nome_docx)

    else:
        print("\nNão foi possível processar o PDF.")